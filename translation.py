import string

import ezodf
import nltk
import openai
from docx import Document
from nltk.tokenize import sent_tokenize, word_tokenize

nltk.download('punkt')

BLUE_TEXT = '\033[34m{}\033[0m'
GREEN_TEXT = '\033[32m{}\033[0m'
RED_TEXT = '\033[31m{}\033[0m'
YELLOW_TEXT = '\033[33m{}\033[0m'

GPT_3_5_TURBO_MAX_TOKENS = 4096


def split_into_sentences(text):
    sentence_list = sent_tokenize(text)
    return sentence_list


def clear_node_content(node):
    for child in node:
        clear_node_content(child)
    node.text = ''
    node.tail = ''


def get_max_tokens(prompt):
    prompt_tokens = len(word_tokenize(prompt))
    max_tokens = GPT_3_5_TURBO_MAX_TOKENS - prompt_tokens
    return max_tokens


def translate_chunk(text_chunk, prompt):
    user_message = f'{prompt} text: {text_chunk}'
    response = openai.chat.completions.create(
        model='gpt-3.5-turbo',
        temperature=0,
        messages=[
            {'role': 'system',
             'content': prompt},
            {'role': 'user',
             'content': user_message}
        ]
    )
    tokens_used = len(word_tokenize(text_chunk))
    print(f"Tokens used: {tokens_used}")
    return response.choices[0].message.content


def translate_node(node, prompt):
    if (isinstance(node, ezodf.text.Paragraph) or
            isinstance(node, ezodf.text.Heading) or
            isinstance(node, ezodf.text.List)):
        inner_text = node.plaintext()
        if inner_text:
            print(BLUE_TEXT.format(f'\n***\nNode: {node}'))
            print(GREEN_TEXT.format('SOURCE_TEXT:'), inner_text)

            translated_text = translate_text(inner_text, prompt)
            print(GREEN_TEXT.format('TRANSLATED_TEXT: '), translated_text)
            clear_node_content(node)
            node.text = translated_text
    elif isinstance(node, ezodf.table.Table):
        for row in node:
            for cell in row:
                for element in cell:
                    translate_node(element, prompt)


def translate_text(text, prompt):
    if not text.strip() or text in string.punctuation:
        return text

    max_tokens = get_max_tokens(prompt)
    openai.api_key = "API_KEY"

    translated_text = ""
    sentences = split_into_sentences(text)
    text_chunk = ""

    for sentence in sentences:
        if len(word_tokenize(text_chunk + sentence)) < max_tokens:
            text_chunk += sentence
        else:
            translated_text += translate_chunk(text_chunk, prompt)
            text_chunk = sentence

    if text_chunk:
        translated_text += translate_chunk(text_chunk, prompt)

    return translated_text


def translate_paragraph(paragraph, prompt):
    text = ''.join(run.text for run in paragraph.runs)
    if text.strip():
        print(BLUE_TEXT.format(f'\n***\nParagraph: {paragraph}'))
        print(GREEN_TEXT.format('SOURCE_TEXT:'), text)

        if paragraph.style.name.startswith('List'):
            list_items = text.split('\n')
            translated_text = ""
            for item in list_items:
                translated_text += translate_text(item, prompt) + "\n"
        else:
            translated_text = translate_text(text, prompt)

        print(GREEN_TEXT.format('TRANSLATED_TEXT: '), translated_text)
        for run in paragraph.runs:
            run.text = ''
        paragraph.runs[0].text = translated_text


def translate_docx_file(docx_file, prompt):
    doc = Document(docx_file)
    for i, paragraph in enumerate(doc.paragraphs):
        translate_paragraph(paragraph, prompt)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    translate_paragraph(paragraph, prompt)

    doc.save(docx_file)
